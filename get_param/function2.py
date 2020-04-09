#!/usr/bin/env python
# Short demo kmpfit (04-03-2012)


import numpy as np
import fitter_3ch
import json
from pylab import *
from scipy.optimize import minimize
from collections import defaultdict
from kapteyn import kmpfit
import win32com.client
import matplotlib
import matplotlib.pyplot as plt
import plotly as py
import plotly.express as px
import plotly.graph_objs as go
from docx import Document
import docx

def model(p, x): #p - наилучшие параметры
   a,b,c,d = p
   return a + b*x + c*x*x + d*x*x*x






def residuals(p, data):  # Residuals function needed by kmpfit
   x, y, er = data           # Data arrays is a tuple given by programmer
   #er = er/y
   #er_sr = np.sum(er)/len(er)
   a, b, c, d = p              # Parameters which are adjusted by kmpfit
   return (y-(a+(b*x)+c*(x**2)+d*(x**3)))/er







def confidence_band(x, dfdp, confprob, fitobj, model, err, abswei):
   from scipy.stats import t
   # Given the confidence probability confprob = 100(1-alpha)
   # we derive for alpha: alpha = 1 - confprob


   alpha = 1.0 - confprob
   prb = 1.0 - alpha/2
   tval = t.ppf(prb, fitobj.dof)

   C = fitobj.covar
   n = len(fitobj.params)              # Number of parameters from covariance matrix
   p = fitobj.params
   N = len(x)
   if abswei:
      covscale = 1.0
   else:
      covscale = fitobj.rchi2_min
   df2 = np.zeros(N)
   for j in range(n):
      for k in range(n):
         df2 += dfdp[j]*dfdp[k]*C[j,k]
   df = np.sqrt(err*err+covscale*df2)
   y = model(p, x)
   delta = tval * df
   upperband = y + delta
   lowerband = y - delta

   return y, upperband, lowerband, delta





def build(data, LU, eror_, ERRS, N,  confprob, d, Xname = '  ', Yname = '   ', A=0, B=1):
	#сортировка всплесков по заданной велечине (по яркости или по жескости) и построение 
	#зависимостей и апроксимация для каждого интервала по яркости или жескости  
	
	#data - словарь с R и cos(Tetta)
	#LU - словарь с яркостью или жескостью, по которой распределяем всплески и cos(Tetta)
	#ERRS - словарь с cos(Tetta) и ID всплесков (нужен для поиска ID всплесков, выпадающих из апроксимации)
	#N - массив с номерами интервалов по яркоти, графики которых нужно вывести 
	#confprob - доверительный интервал 
	#d - колличество интервалов, на которых разбиваем всплески
	#Xname и Yname - подписи к графикам 
	a=[]

	l = np.sort(list(LU.keys()))
	
	for j in range(d):
	
		i = 0
		R1 = dict()

		if j == 0:
			S = 0
		else:
			S = F 
		if j == d-1:
			F = len(l)
		else:
			F = (j+1)*round(len(l)/d)

		for i in range(S, F, 1):
			R1[LU[l[i]]] = data[LU[l[i]]]
	
		a.append(R1)
		#print(a[j],type(a[j]))

	graf(a, eror_, ERRS, N, confprob, Xname, Yname, A, B)

	

AW = []
CHi_ = [[ [], [] ]]

I=0

def graf(a, eror_, ERRS, N, confprob, Xname = '  ', Yname = '   ', A=0, B=1):
	#построение зависимостей и апроксимация для каждого интервала по яркости или жескости
	#вынесен в отдельную функцию, что бы можно было бы миновать этап разбиения и получить
	#все всплески на одном графике
	
	#a - массив со словарями, содержащими R и cos(Tetta), распределенными по жескостям или яркостям, 
	#получен из функции build или словарь с R и cos(Tetta) если хочу миновать этап разбиения на интервалы 
	#и получить все всплески на одном графике
	
	#eror_ - словарь с погрешностью R и значениями cos(Tetta)
	#ERRS - словарь с cos(Tetta) и ID всплесков (нужен для поиска ID всплесков, выпадающих из апроксимации и для подписей к графикам)
	#N - массив с номерами интервалов по яркоти, графики которых нужно вывести 
	#confprob - 
	#Xname и Yname - подписи к графикам 
	#ID - необязательный параметр, если задать ID=1, то каждая точка на графике будет подписана ее ID

	CHi1 = []
	PARAM = []
	
	ii = 1

	if type(a) is not list: 
		a = [a]
	
	#print(len(a[0]))
	for j in N:
		
		Data = a[j-1]

		X = np.array(list(Data.keys()))
		Y = np.array(list(Data.values()))
		dfdp = [1, X, X**2, X**3]
		err = np.zeros(len(X))
		err_1 = np.zeros(len(X))
		for i in range(len(X)):
			err[i] = eror_[X[i]]
			err_1[i] = eror_[X[i]]/Y[i]

		print(len(X))
		print(Y[0], Y[len(Y)-1])

		paramsinitial = [-0.0149, 0.00669, -0.0000889, 0.00000158]
		fitobj = kmpfit.Fitter(residuals=residuals, data=(X,Y,err), maxiter = 5000)
		fitobj.fit(params0=paramsinitial)

		erR=np.sqrt((np.sum((Y - (fitobj.params[0] + fitobj.params[1]*X + fitobj.params[2]*X**2 + fitobj.params[3]*X**3))**2))/len(X))
		eRR = erR*np.ones(len(X))


		yhat, upperband, lowerband, delta = confidence_band(X, dfdp, confprob, fitobj, model,eRR, abswei=False )

		

		title_name = "{} {} {}  {} {} {}".format(
			Yname, '_', j, '(', Xname, ')')
			
		title_name = ''.join(title_name.split())

		#параметры выводятся начиная с параметра при старшей степени
		Title_name = "{}\n {} {} {} {} {} {}\n {} {} {} {} {} {} {} {}".format(
			title_name, 'χ2 =', round(fitobj.chi2_min/len(X), 3), 'dof=' ,len(X), 'delta=',round(np.sum(delta)/len(delta),3), 'a=', round(fitobj.params[3],3), 'b=', round(fitobj.params[2],3), 'c=', round(fitobj.params[1],3),'d=',round(fitobj.params[0],3))

		#Title_name = "{}\n  {}\n {} {} {} {} {} {} {} {} {} ".format(
		#	title_name, 'χ2=', round(fitobj.chi2_min/len(X), 3), 'd=', len(X), 'Δ=',  round(np.sum(delta)/len(delta),2), round(fitobj.params[3],2), round(fitobj.params[2],2),  round(fitobj.params[1],2),round(fitobj.params[0],2))


		"""
		wr = "{} {} {} {} {}   {}   {}   {}   {}\n".format(
			Yname + '_' + str(j), round(fitobj.chi2_min/(len(X)-1), 3),  j, round(np.sum(delta)/len(delta),10) ,len(X)-1, round(fitobj.params[3], 10), round(fitobj.params[2], 10), round(fitobj.params[1], 10), round(fitobj.params[0], 10) ) 

		wr2 = "{:15.3f} {:15.3f} {:15.0f} {:15.10f} {:15.10f} {:15.10f} {:15.10f} {:15s}\n".format(
			 round(fitobj.chi2_min/(len(X)-1), 3), round(np.sum(delta)/len(delta),10) ,len(X)-1, round(fitobj.params[3], 10), round(fitobj.params[2], 10), round(fitobj.params[1], 10), round(fitobj.params[0], 10),Yname + '_' + str(j)) 
		

		if 'S1' in Yname:
			parametrS1 = open('P_S1.txt', 'a')
			parametrS1.write(str(wr))
			parametrS1.close()
			parametrS1_2 = open('parametrs_S1.txt', 'a')
			parametrS1_2.write(str(wr2))
			parametrS1_2.close()		

		else:
			parametrS2 = open('P_S2.txt', 'a')
			parametrS2.write(str(wr))
			parametrS2.close()
			parametrS2_2 = open('parametrs_S2.txt', 'a')
			parametrS2_2.write(str(wr2))
			parametrS2_2.close()		

		"""

		title_name = "{} {} {}  {} {} {} {}".format(
			Yname, '_', j, '(', Xname, ')','.html')

		title_name = ''.join(title_name.split())





		"""
		X_t=[]
		Y_t=[]
		err_t=[]

		
		for i,j in ERRS.items():
			if j == 1205 or j == 1422 or j == 1428 or j == 4122 or  j == 3757 or j == 4504 or j == 4740  or j == 2935 or j == 3357: #or j == 2982
			#if j == 3370:
				X_t.append(i)
				Y_t.append(Data[i])
				err_t.append(eror_[i])
				list(X).remove(i)
				list(Y).remove(Data[i])	
				list(err).remove(eror_[i])
		
		ID_t=[]
		for i in X_t:
			ID_t.append(ERRS[i]) 

		"""

		ID=[]
		for i in X:
			ID.append(ERRS[i]) 

	
			
		fig=go.Scatter(x =X, y=Y, marker={'color': 'black',  'size': 4},
		mode = 'markers', text = ID, error_y = dict(type = 'data',width = 0, color = 'red', thickness=1, array = err))
		#fig4=go.Scatter(x =X_t, y=Y_t, marker={'color': 'red',  'size': 10},
		#mode = 'markers', text = ID_t, error_y = dict(type = 'data',width = 0, color = 'red', thickness=1, array = err_t))
		

		t = np.linspace(A, B, 100)

		fig1=go.Scatter(x=t, y=fitobj.params[0] + fitobj.params[1]*t + fitobj.params[2]*t**2 + fitobj.params[3]*t**3,  mode = 'lines')
		fig2=go.Scatter(x=t, y=fitobj.params[0] + fitobj.params[1]*t + fitobj.params[2]*t**2 + fitobj.params[3]*t**3 + np.sum(delta)/len(delta), marker={'color': 'red'}, mode = 'lines')
		fig3=go.Scatter(x=t, y=fitobj.params[0] + fitobj.params[1]*t + fitobj.params[2]*t**2 + fitobj.params[3]*t**3 - np.sum(delta)/len(delta), marker={'color': 'red'}, mode = 'lines')
		

		data = [fig1, fig2, fig3, fig]
		layout = {'title': Title_name, 'xaxis': {'title': Xname, 'range' : [A-0.1,B+0.1] }, 'yaxis': {'title': Yname, 'range': [-0.2, 1.2] }}
		



		#fig.write_image("graf/FFF.WebP")
		#fig.show()

		FIG = go.Figure(data=data, layout=layout)
		
		py.offline.plot(FIG, filename='Incident Angle/' + title_name)

		#FIG.write_image("F.jpg")

		#doc = Document('picture.docx')

		#doc = Document()

		#doc.add_paragraph('Это первый абзац')
		#doc.add_picture('F.jpg', width = docx.shared.Cm(10))

		#doc.save('picture.docx')






		#print('chi2 for', Yname, j, ':', ((np.sum(err))**2)*fitobj.chi2_min/(len(X))**3)
		
		print('chi2 for', Yname, j, ':', fitobj.chi2_min/len(X))
		print(round(fitobj.params[3], 3),  round(fitobj.xerror[3],3), round(fitobj.params[2],3),  round(fitobj.xerror[2],3), round(fitobj.params[1],3), round(fitobj.xerror[1],3), round(fitobj.params[0],3),  round(fitobj.xerror[0],3), round(np.sum(delta)/len(delta),3), round(np.sum(delta)/len(delta),3), round(fitobj.chi2_min/len(X),3),'\n')
		


		#print('chi2 for', Yname, j, ':', fitobj.chi2_min/len(X))
		#print('a=',fitobj.params[3], '+-', fitobj.xerror[3],'\n','b=',fitobj.params[2], '+-', fitobj.xerror[2],'\n','c=',fitobj.params[1], '+-', fitobj.xerror[1],'\n','d=',fitobj.params[0], '+-', fitobj.xerror[0],'\n')
		#print(Yname, j, ':')
		#print(len(X))
		#print(fitobj.params)
		#print(np.sum(delta)/len(delta))
		#print(fitobj.params[0]+np.sum(delta)/len(delta))
		#print(fitobj.params[0]-np.sum(delta)/len(delta))
		#print('chi2 for', Yname, ii, ':', fitobj.chi2_min*(np.sum(err)**2/len(err)**2)/len(X))
		
	 
		ii = ii + 1 
			
		#цикл отыскивает всплески, сильно выпадающие из аппроксимации (см 231 строчку) 	
		for i in range(len(X)):

			if (Y[i] - (fitobj.params[0] + fitobj.params[1]*X[i] + fitobj.params[2]*X[i]**2 + fitobj.params[3]*X[i]**3))**2 > (0.2)**2:
				if ERRS[X[i]] not in AW:
					AW.append(ERRS[X[i]])

		

		#CHi1.append(round( ((np.sum(err))**2)*fitobj.chi2_min/(len(X))**3 ,4))
		CHi1.append(round(fitobj.chi2_min/(len(X)) ,3))
		PARAM.append([round(fitobj.params[3],2), round(fitobj.params[2],2), round(fitobj.params[1],2), round(fitobj.params[0],2), len(X)])

	global I

	if len(CHi_[I])<=12:
		CHi_[I][0].append(CHi1)
		CHi_[I][1].append(PARAM)
	else:
		I=I+1
		CHi_[I][0].append(CHi1)
		CHi_[I][1].append(PARAM)
		 



def print_chi(*Del):
	#печатает в exel файл значения всех chi2, запускается полсе запуска построения всех для комбинаций каналов
	#вначале детектора S1, потом детектора S2 для любого распредления по интервалам
	#в таблицы вводит данные в  той последовательности, которая представлена в отчете  
	#если выдает ошибку, то нужно перед запуском скрипта открыть exel файлы
	#DEl - необязательный параметр - каналы, которые не будут рассматривать как дающие лучшие значение chi2 (например G1; G1, G3)
	Excel = win32com.client.Dispatch("Excel.Application")

	wb = Excel.Workbooks.Open(u'C:\\Users\\Gildor\\python\\chi2.xlsx')
	Slon = wb.ActiveSheet

	chi2_best_S1 = Excel.Workbooks.Open(u'C:\\Users\\Gildor\\python\\chi2_best_S1.xlsx')
	sheet1 = chi2_best_S1.ActiveSheet
	
	chi2_best_S2 = Excel.Workbooks.Open(u'C:\\Users\\Gildor\\python\\chi2_best_S2.xlsx')
	sheet2 = chi2_best_S2.ActiveSheet
	

	for CHi_FULL in CHi_:

		CHi = CHi_FULL[0]
		PAR = CHi_FULL[1]

		for i in range(len(CHi[0])):
			dic1 = dict()
			par1_1 = dict()
			par1_2 = dict()
			par1_3 = dict()
			par1_4 = dict()
			len1_1 = dict()

			dic2 = dict()
			par2_1 = dict()
			par2_2 = dict()
			par2_3 = dict()
			par2_4 = dict()
			len2_1 = dict()

			i = i+1
			k=0
			for j in range(0,len(CHi),2):
				j = j + 1

				
				Slon.Cells(i,j).value = CHi[k][i-1]

				if (k+1)==1:
					dic1[CHi[k][i-1]] = 'G1' #''.join(str('G',k+1).split()) 
					
					par1_1['G1'] = PAR[k][i-1][0]
					par1_2['G1'] = PAR[k][i-1][1]
					par1_3['G1'] = PAR[k][i-1][2]
					par1_4['G1'] = PAR[k][i-1][3]
					len1_1['G1'] = PAR[k][i-1][4]

				if (k+1)==2:
					dic1[CHi[k][i-1]] = 'G2' #''.join(str('G',k+1).split()) 
					
					par1_1['G2'] = PAR[k][i-1][0]
					par1_2['G2'] = PAR[k][i-1][1]
					par1_3['G2'] = PAR[k][i-1][2]
					par1_4['G2'] = PAR[k][i-1][3]
					len1_1['G2'] = PAR[k][i-1][4]

				if (k+1)==3:
					dic1[CHi[k][i-1]] = 'G3' #''.join(str('G',k+1).split()) 
					
					par1_1['G3'] = PAR[k][i-1][0]
					par1_2['G3'] = PAR[k][i-1][1]
					par1_3['G3'] = PAR[k][i-1][2]
					par1_4['G3'] = PAR[k][i-1][3]
					len1_1['G3'] = PAR[k][i-1][4]

				if (k+1)==4:
					dic1[CHi[k][i-1]] = str('G1G2')
					
					par1_1[str('G1G2')] = PAR[k][i-1][0]
					par1_2[str('G1G2')] = PAR[k][i-1][1]
					par1_3[str('G1G2')] = PAR[k][i-1][2]
					par1_4[str('G1G2')] = PAR[k][i-1][3]
					len1_1[str('G1G2')] = PAR[k][i-1][4]

				if (k+1)==5:
					dic1[CHi[k][i-1]] = str('G2G3')
					
					par1_1[str('G2G3')] = PAR[k][i-1][0]
					par1_2[str('G2G3')] = PAR[k][i-1][1]
					par1_3[str('G2G3')] = PAR[k][i-1][2]
					par1_4[str('G2G3')] = PAR[k][i-1][3]
					len1_1[str('G2G3')] = PAR[k][i-1][4]

				if (k+1)==6:	
					dic1[CHi[k][i-1]] = str('G1G2G3')
					
					par1_1[str('G1G2G3')] = PAR[k][i-1][0]
					par1_2[str('G1G2G3')] = PAR[k][i-1][1]
					par1_3[str('G1G2G3')] = PAR[k][i-1][2]
					par1_4[str('G1G2G3')] = PAR[k][i-1][3]
					len1_1[str('G1G2G3')] = PAR[k][i-1][4]

				k=k+1
			for j in range(1,len(CHi),2):
				j = j + 1

				Slon.Cells(i,j).value = CHi[k][i-1]


				if (k+1-6)==1:
					dic2[CHi[k][i-1]] = 'G1' #''.join(str('G',k+1-6).split())

					par2_1['G1'] = PAR[k][i-1][0]
					par2_2['G1'] = PAR[k][i-1][1]
					par2_3['G1'] = PAR[k][i-1][2]
					par2_4['G1'] = PAR[k][i-1][3]
					len2_1['G1'] = PAR[k][i-1][4]

				if (k+1-6)==2:
					dic2[CHi[k][i-1]] = 'G2' #''.join(str('G',k+1-6).split())
					
					par2_1['G2'] = PAR[k][i-1][0]
					par2_2['G2'] = PAR[k][i-1][1]
					par2_3['G2'] = PAR[k][i-1][2]
					par2_4['G2'] = PAR[k][i-1][3]
					len2_1['G2'] = PAR[k][i-1][4]
				
				if (k+1-6)==3:
					dic2[CHi[k][i-1]] = 'G3' #''.join(str('G',k+1-6).split())	
					
					par2_1['G3'] = PAR[k][i-1][0]
					par2_2['G3'] = PAR[k][i-1][1]
					par2_3['G3'] = PAR[k][i-1][2]
					par2_4['G3'] = PAR[k][i-1][3]
					len2_1['G3'] = PAR[k][i-1][4]

				if (k+1-6)==4:
					dic2[CHi[k][i-1]] = str('G1G2')
					
					par2_1[str('G1G2')] = PAR[k][i-1][0]
					par2_2[str('G1G2')] = PAR[k][i-1][1]
					par2_3[str('G1G2')] = PAR[k][i-1][2]
					par2_4[str('G1G2')] = PAR[k][i-1][3]
					len2_1[str('G1G2')] = PAR[k][i-1][4]

				if (k+1-6)==5:
					dic2[CHi[k][i-1]] = str('G2G3')
					
					par2_1[str('G2G3')] = PAR[k][i-1][0]
					par2_2[str('G2G3')] = PAR[k][i-1][1]
					par2_3[str('G2G3')] = PAR[k][i-1][2]
					par2_4[str('G2G3')] = PAR[k][i-1][3]
					len2_1[str('G2G3')] = PAR[k][i-1][4]

				if (k+1-6)==6:	
					dic2[CHi[k][i-1]] = str('G1G2G3')
					
					par2_1[str('G1G2G3')] = PAR[k][i-1][0]
					par2_2[str('G1G2G3')] = PAR[k][i-1][1]
					par2_3[str('G1G2G3')] = PAR[k][i-1][2]
					par2_4[str('G1G2G3')] = PAR[k][i-1][3]
					len2_1[str('G1G2G3')] = PAR[k][i-1][4]
					
				k=k+1

				for DEL in Del:	
					s=[]
					while DEL in dic1.values():
						for i1, j in dic1.items():
							if j == DEL:
								s.append(i1)
						for i1 in s:
							del dic1[i1]
					s = []		
					while DEL in dic2.values():
						for i1, j in dic2.items():
							if j == DEL:
								s.append(i1)
						for i1 in s:
							del dic2[i1]
			
			list1 = list(np.sort(list(dic1.keys())))
			F1 = list1.pop(0)
			
			#d = PAR[i][CHi[i].index(F1)] 


			#g = [PAR[CHi.index(F1.all())][0], PAR[CHi.index(F1.all())][1], PAR[CHi.index(F1.all())][2], PAR[CHi.index(F1.all())][3]]
			#hhhh = ', '.join(g) 

			
			
			sheet1.Cells(i,1).value = dic1[F1]

			sheet1.Cells(i,2).value = len1_1[dic1[F1]]

			sheet1.Cells(i,3).value = F1

			n = [par1_1[dic1[F1]], par1_2[dic1[F1]], par1_3[dic1[F1]], par1_4[dic1[F1]]]
			
			sheet1.Cells(i,4).value = ', '.join(str(z) for z in n)

			
			F2 = list1.pop(0)

			sheet1.Cells(i,5).value = dic1[F2]

			sheet1.Cells(i,6).value = len1_1[dic1[F2]]

			sheet1.Cells(i,7).value = F2

			n = [par1_1[dic1[F2]], par1_2[dic1[F2]], par1_3[dic1[F2]], par1_4[dic1[F2]]]


			sheet1.Cells(i,8).value = ', '.join(str(z) for z in n)

			list2 = list(np.sort(list(dic2.keys())))
			F3 = list2.pop(0)
			
			sheet2.Cells(i,1).value = dic2[F3]

			sheet2.Cells(i,2).value = len2_1[dic2[F3]]

			sheet2.Cells(i,3).value = F3

			n = [par2_1[dic2[F3]], par2_2[dic2[F3]], par2_3[dic2[F3]], par2_4[dic2[F3]]]

			sheet2.Cells(i,4).value = ', '.join(str(z) for z in n)
			F4 = list2.pop(0)
		
			sheet2.Cells(i,5).value = dic2[F4]

			sheet2.Cells(i,6).value = len2_1[dic2[F4]]

			sheet2.Cells(i,7).value = F4

			n = [par2_1[dic2[F4]], par2_2[dic2[F4]], par2_3[dic2[F4]], par2_4[dic2[F4]]]

			sheet2.Cells(i,8).value = ', '.join(str(z) for z in n)
		
		wb.Close()
		chi2_best_S1.Close()
		chi2_best_S2.Close()
		
		Excel.Quit()



	def print_er(LUM=0, h=3):
		#печатает всплески, из построенных этим запуском скрипта графиков, сильно (см массив AW из функции graf и 231 строчку) отклоняющиеся от аппроксимайции 
		#(Необязательная переменная LUM - словарь с яркостью и ID. 
		#Если задана, то в файле с результатами рядом с ID в части 1/h самых неярких всплесков пишет "DA") 
		ER = open('er.txt', 'w')

		if LUM != 0: 
			LU = sort(list(LUM.keys()))
			ID_mzgkie = [] 

			for i in range(int(len(LU)/h)):           
				ID_mzgkie.append(LUM[LU[i]])

				

			for i in range(len(AW)):
				if AW[i] in ID_mzgkie:
					aw = "{:10.0f} {:1s} {:2s}\n".format(
						AW[i], ',', 'DA')
					ER.write(str(aw))
				else:
					aw = "{:10.0f} {:1s}\n".format(
						AW[i], ',')
					ER.write(str(aw))

		else:
			for i in range(len(AW)):
				aw = "{:10.0f} {:1s}\n".format(
						AW[i], ',')
				ER.write(str(aw))

		ER.close()











kw_cuts = np.array([13.125, 50.0, 200.0, 750.0])

def calc_cal_HR_fit(counts, counts_err, fcal):
	#функция для перевода из исходых в номинальные
 
	cuts_cur = kw_cuts * fcal
	cuts_nom = kw_cuts
 
	fitter = fitter_3ch.Fitter(fitter_3ch.model('CPL'), cuts_cur, counts, counts_err)
	result, chi2 = fitter.fit()

       
	fitted_counts = fitter.model.int_func(result, cuts_nom[0:3], cuts_nom[1:4])
	#print ("inp counts: "), counts
	#print ("out counts: "), fitted_counts
 
	return fitted_counts, result, chi2


#CHI.close()
